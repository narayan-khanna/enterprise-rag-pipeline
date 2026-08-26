"""
Auto Loader
===========
Detects file type from extension (or URL) and loads content into
LangChain Document objects. Supported formats:

    .pdf   -> PyPDFLoader   (one Document per page)
    .docx  -> UnstructuredWordDocumentLoader
    .txt   -> TextLoader
    .md    -> TextLoader
    .csv   -> CSVLoader     (one Document per row)
    .json  -> JSONLoader
    .jsonl -> line-by-line JSON
    URL    -> WebBaseLoader

Usage:
    from rag.loaders.auto_loader import AutoLoader

    loader = AutoLoader()
    docs = loader.load("report.pdf")
    docs = loader.load(["doc1.pdf", "doc2.txt"])
    docs = loader.load("https://example.com/article")
    docs = loader.load_text("Some plain text...", source="my_notes")
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List, Union

from langchain.schema import Document


class AutoLoader:
    """File-type-aware document loader."""

    _EXTENSION_MAP = {
        ".pdf": "_load_pdf",
        ".txt": "_load_text",
        ".md": "_load_text",
        ".csv": "_load_csv",
        ".json": "_load_json",
        ".jsonl": "_load_jsonl",
        ".docx": "_load_docx",
        ".doc": "_load_docx",
        ".html": "_load_html",
        ".htm": "_load_html",
    }

    def load(self, source: Union[str, Path, List]) -> List[Document]:
        if isinstance(source, list):
            docs: List[Document] = []
            for s in source:
                docs.extend(self.load(s))
            return docs

        source_str = str(source)

        if source_str.startswith("http://") or source_str.startswith("https://"):
            return self._load_url(source_str)

        path = Path(source_str)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source_str}")

        if path.is_dir():
            return self._load_directory(path)

        method_name = self._EXTENSION_MAP.get(path.suffix.lower())
        if method_name is None:
            return self._load_text(str(path))

        return getattr(self, method_name)(str(path))

    def load_text(self, text: str, source: str = "inline") -> List[Document]:
        return [Document(page_content=text, metadata={"source": source})]

    def _load_pdf(self, path: str) -> List[Document]:
        from langchain.document_loaders import PyPDFLoader
        return PyPDFLoader(path).load()

    def _load_text(self, path: str) -> List[Document]:
        from langchain.document_loaders import TextLoader
        return TextLoader(path, encoding="utf-8").load()

    def _load_csv(self, path: str) -> List[Document]:
        from langchain.document_loaders import CSVLoader
        return CSVLoader(path).load()

    def _load_json(self, path: str) -> List[Document]:
        from langchain.document_loaders import JSONLoader
        try:
            return JSONLoader(path, jq_schema=".[]", text_content=False).load()
        except Exception:
            content = Path(path).read_text(encoding="utf-8")
            return [Document(page_content=content, metadata={"source": path})]

    def _load_jsonl(self, path: str) -> List[Document]:
        docs: List[Document] = []
        with open(path, encoding="utf-8") as f:
            for i, line in enumerate(f):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                    content = json.dumps(obj, ensure_ascii=False)
                except json.JSONDecodeError:
                    content = line
                docs.append(Document(page_content=content, metadata={"source": path, "line": i}))
        return docs

    def _load_docx(self, path: str) -> List[Document]:
        try:
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
            return UnstructuredWordDocumentLoader(path).load()
        except ImportError:
            import docx
            doc = docx.Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            return [Document(page_content=text, metadata={"source": path})]

    def _load_html(self, path: str) -> List[Document]:
        try:
            from bs4 import BeautifulSoup
            html = Path(path).read_text(encoding="utf-8")
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(separator="\n")
            return [Document(page_content=text, metadata={"source": path})]
        except ImportError:
            return self._load_text(path)

    def _load_url(self, url: str) -> List[Document]:
        from langchain.document_loaders import WebBaseLoader
        return WebBaseLoader(url).load()

    def _load_directory(self, path: Path) -> List[Document]:
        docs: List[Document] = []
        for child in sorted(path.rglob("*")):
            if child.is_file() and child.suffix.lower() in self._EXTENSION_MAP:
                try:
                    docs.extend(self.load(str(child)))
                except Exception as e:
                    print(f"[AutoLoader] Skipping {child}: {e}")
        return docs
